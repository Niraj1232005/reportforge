import base64
import re
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional
from urllib.parse import unquote

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import html as lxml_html


IMAGE_PATTERN = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
IMAGE_REF_PREFIX = "rf-image://"
HTML_TAG_PATTERN = re.compile(r"<[a-zA-Z][^>]*>")
PX_PATTERN = re.compile(r"([0-9]+(?:\.[0-9]+)?)px")
PT_PATTERN = re.compile(r"([0-9]+(?:\.[0-9]+)?)pt")
PERCENT_PATTERN = re.compile(r"([0-9]+(?:\.[0-9]+)?)%")
CITATION_PATTERN = re.compile(r"\[@([a-zA-Z0-9_-]+)\]")
FOOTNOTE_PATTERN = re.compile(r"\[fn:([a-zA-Z0-9_-]+)\]")

DEFAULT_DOCUMENT_SETTINGS: Dict[str, Any] = {
    "fontFamily": "Times New Roman",
    "bodyFontSize": 12,
    "heading1Size": 18,
    "heading2Size": 16,
    "heading3Size": 14,
    "paragraphAlign": "justify",
    "lineSpacing": 1.5,
    "marginTopIn": 1,
    "marginBottomIn": 1,
    "marginLeftIn": 1,
    "marginRightIn": 1,
    "pageBreakAfterHeading1": True,
}

DEFAULT_DOCUMENT_STRUCTURE: Dict[str, Any] = {
    "showCoverPage": False,
    "showTableOfContents": False,
}


def _normalize_document_settings(value: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    source = value if isinstance(value, dict) else {}

    def _number(name: str, fallback: float, minimum: float, maximum: float) -> float:
        try:
            return max(minimum, min(maximum, float(source.get(name, fallback))))
        except (TypeError, ValueError):
            return fallback

    paragraph_align = str(source.get("paragraphAlign") or DEFAULT_DOCUMENT_SETTINGS["paragraphAlign"]).strip().lower()
    if paragraph_align not in {"left", "right", "center", "justify"}:
        paragraph_align = DEFAULT_DOCUMENT_SETTINGS["paragraphAlign"]

    font_family = str(source.get("fontFamily") or DEFAULT_DOCUMENT_SETTINGS["fontFamily"]).strip() or DEFAULT_DOCUMENT_SETTINGS["fontFamily"]

    return {
        "fontFamily": font_family,
        "bodyFontSize": _number("bodyFontSize", 12, 10, 16),
        "heading1Size": _number("heading1Size", 18, 14, 28),
        "heading2Size": _number("heading2Size", 16, 13, 24),
        "heading3Size": _number("heading3Size", 14, 12, 20),
        "paragraphAlign": paragraph_align,
        "lineSpacing": _number("lineSpacing", 1.5, 1.0, 2.0),
        "marginTopIn": _number("marginTopIn", 1, 0.5, 2.0),
        "marginBottomIn": _number("marginBottomIn", 1, 0.5, 2.0),
        "marginLeftIn": _number("marginLeftIn", 1, 0.5, 2.0),
        "marginRightIn": _number("marginRightIn", 1, 0.5, 2.0),
        "pageBreakAfterHeading1": bool(
            source.get("pageBreakAfterHeading1", True)
        ),
    }


def _normalize_document_structure(value: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    source = value if isinstance(value, dict) else {}
    return {
        "showCoverPage": bool(source.get("showCoverPage")),
        "showTableOfContents": bool(source.get("showTableOfContents")),
    }


def _apply_style_font(style, font_name: str, size_pt: int, bold: Optional[bool] = None):
    style.font.name = font_name
    style.font.size = Pt(size_pt)

    if bold is not None:
        style.font.bold = bold

    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _configure_document(doc: Document, settings: Dict[str, Any]):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(float(settings["marginTopIn"]))
    section.bottom_margin = Inches(float(settings["marginBottomIn"]))
    section.left_margin = Inches(float(settings["marginLeftIn"]))
    section.right_margin = Inches(float(settings["marginRightIn"]))

    font_name = str(settings["fontFamily"])
    _apply_style_font(doc.styles["Normal"], font_name, size_pt=int(settings["bodyFontSize"]), bold=False)
    _apply_style_font(doc.styles["Title"], font_name, size_pt=int(settings["heading1Size"]) + 4, bold=True)
    _apply_style_font(doc.styles["Heading 1"], font_name, size_pt=int(settings["heading1Size"]), bold=True)
    _apply_style_font(doc.styles["Heading 2"], font_name, size_pt=int(settings["heading2Size"]), bold=True)
    _apply_style_font(doc.styles["Heading 3"], font_name, size_pt=int(settings["heading3Size"]), bold=True)


def _normalized_text(value: Any, fallback: str) -> str:
    if value is None:
        return fallback

    text = str(value).strip()
    return text if text else fallback


def _looks_like_html(content: str) -> bool:
    return bool(HTML_TAG_PATTERN.search(content))


def _parse_style_map(style_value: str) -> Dict[str, str]:
    parsed: Dict[str, str] = {}

    for declaration in style_value.split(";"):
        if ":" not in declaration:
            continue

        key, value = declaration.split(":", 1)
        key = key.strip().lower()
        value = value.strip()
        if key and value:
            parsed[key] = value

    return parsed


def _set_run_font_name(run, font_name: str):
    run.font.name = font_name
    r = run._r
    rpr = r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _apply_inline_run_style(run, style: Dict[str, Any]):
    run.bold = style.get("bold")
    run.italic = style.get("italic")
    run.underline = style.get("underline")
    run.font.highlight_color = style.get("highlight")

    if isinstance(style.get("font_name"), str):
        _set_run_font_name(run, style["font_name"])

    if isinstance(style.get("font_size_pt"), (int, float)):
        run.font.size = Pt(float(style["font_size_pt"]))


def _parse_hex_color(value: str) -> Optional[tuple]:
    normalized = value.strip().lower()
    if not normalized.startswith("#"):
        return None

    raw = normalized[1:]
    if len(raw) == 3:
        raw = "".join([char * 2 for char in raw])

    if len(raw) != 6 or not re.fullmatch(r"[0-9a-f]{6}", raw):
        return None

    return tuple(int(raw[index : index + 2], 16) for index in (0, 2, 4))


def _parse_rgb_color(value: str) -> Optional[tuple]:
    match = re.fullmatch(
        r"rgb\(\s*([0-9]{1,3})\s*,\s*([0-9]{1,3})\s*,\s*([0-9]{1,3})\s*\)",
        value.strip().lower(),
    )
    if not match:
        return None

    channels = tuple(int(match.group(index)) for index in (1, 2, 3))
    if any(channel < 0 or channel > 255 for channel in channels):
        return None

    return channels


def _parse_highlight_color(value: str) -> Optional[WD_COLOR_INDEX]:
    normalized = (value or "").strip().lower()
    if not normalized:
        return None

    named_lookup: Dict[str, WD_COLOR_INDEX] = {
        "yellow": WD_COLOR_INDEX.YELLOW,
        "green": WD_COLOR_INDEX.BRIGHT_GREEN,
        "cyan": WD_COLOR_INDEX.TURQUOISE,
        "magenta": WD_COLOR_INDEX.PINK,
        "red": WD_COLOR_INDEX.RED,
        "blue": WD_COLOR_INDEX.BLUE,
        "darkblue": WD_COLOR_INDEX.DARK_BLUE,
        "navy": WD_COLOR_INDEX.DARK_BLUE,
        "darkred": WD_COLOR_INDEX.DARK_RED,
        "teal": WD_COLOR_INDEX.TEAL,
        "gray": WD_COLOR_INDEX.GRAY_25,
        "grey": WD_COLOR_INDEX.GRAY_25,
    }
    if normalized in named_lookup:
        return named_lookup[normalized]

    rgb_value = _parse_hex_color(normalized) or _parse_rgb_color(normalized)
    if rgb_value is None:
        return None

    red, green, blue = rgb_value
    if red >= 220 and green >= 200 and blue <= 210:
        return WD_COLOR_INDEX.YELLOW
    if green >= 220 and red <= 180 and blue <= 180:
        return WD_COLOR_INDEX.BRIGHT_GREEN
    if blue >= 220 and red <= 180 and green <= 180:
        return WD_COLOR_INDEX.BLUE
    if red >= 220 and blue >= 220 and green <= 180:
        return WD_COLOR_INDEX.PINK
    if red >= 220 and green <= 180 and blue <= 180:
        return WD_COLOR_INDEX.RED
    if red >= 200 and green >= 200 and blue >= 200:
        return WD_COLOR_INDEX.GRAY_25

    target_palette = {
        WD_COLOR_INDEX.YELLOW: (255, 255, 0),
        WD_COLOR_INDEX.BRIGHT_GREEN: (0, 255, 0),
        WD_COLOR_INDEX.TURQUOISE: (0, 255, 255),
        WD_COLOR_INDEX.PINK: (255, 0, 255),
        WD_COLOR_INDEX.RED: (255, 0, 0),
        WD_COLOR_INDEX.BLUE: (0, 0, 255),
        WD_COLOR_INDEX.DARK_BLUE: (0, 0, 128),
        WD_COLOR_INDEX.DARK_RED: (128, 0, 0),
        WD_COLOR_INDEX.TEAL: (0, 128, 128),
        WD_COLOR_INDEX.GRAY_25: (160, 160, 160),
    }

    def color_distance(color_a: tuple, color_b: tuple) -> int:
        return (
            (color_a[0] - color_b[0]) ** 2
            + (color_a[1] - color_b[1]) ** 2
            + (color_a[2] - color_b[2]) ** 2
        )

    return min(target_palette, key=lambda item: color_distance(rgb_value, target_palette[item]))


def _parse_font_size_pt(style_value: Optional[str]) -> Optional[float]:
    if not style_value:
        return None

    px_match = PX_PATTERN.search(style_value)
    if px_match:
        px_value = float(px_match.group(1))
        return px_value * 0.75

    pt_match = PT_PATTERN.search(style_value)
    if pt_match:
        return float(pt_match.group(1))

    return None


def _parse_alignment(value: str) -> Optional[int]:
    normalized = value.lower().strip()

    if normalized in {"left", "start"}:
        return WD_ALIGN_PARAGRAPH.LEFT
    if normalized in {"center", "centre"}:
        return WD_ALIGN_PARAGRAPH.CENTER
    if normalized in {"right", "end"}:
        return WD_ALIGN_PARAGRAPH.RIGHT
    if normalized in {"justify", "justified"}:
        return WD_ALIGN_PARAGRAPH.JUSTIFY

    return None


def _extract_alignment(element, style_map: Optional[Dict[str, str]] = None) -> Optional[int]:
    styles = style_map if style_map is not None else _parse_style_map(element.get("style", ""))
    text_align = styles.get("text-align") or element.get("align")

    if not text_align:
        return None

    return _parse_alignment(text_align)


def _parse_width_inches(value: str, max_width_inches: float) -> Optional[float]:
    if not value:
        return None

    percent_match = PERCENT_PATTERN.search(value)
    if percent_match:
        percent = float(percent_match.group(1))
        return max_width_inches * max(0.1, min(1.0, percent / 100.0))

    px_match = PX_PATTERN.search(value)
    if px_match:
        px_value = float(px_match.group(1))
        return px_value / 96.0

    pt_match = PT_PATTERN.search(value)
    if pt_match:
        pt_value = float(pt_match.group(1))
        return pt_value / 72.0

    try:
        numeric_value = float(value)
        return numeric_value / 96.0
    except (TypeError, ValueError):
        return None


def _clamp_image_width(width: Optional[float], max_width_inches: float) -> float:
    if width is None:
        return max_width_inches

    return max(1.25, min(max_width_inches, width))


def _paragraph_blocks(text: str) -> Iterable[Dict[str, str]]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    chunks = [chunk.strip() for chunk in re.split(r"\n{2,}", normalized) if chunk.strip()]
    return [{"type": "paragraph", "text": chunk} for chunk in chunks]


def _content_blocks(content: str) -> List[Dict[str, str]]:
    text = str(content or "")
    blocks: List[Dict[str, str]] = []
    cursor = 0

    for match in IMAGE_PATTERN.finditer(text):
        start = match.start()
        end = match.end()
        alt = (match.group(1) or "Figure").strip() or "Figure"
        src = (match.group(2) or "").strip()

        blocks.extend(_paragraph_blocks(text[cursor:start]))

        if src:
            blocks.append({"type": "image", "source": src, "alt": alt})

        cursor = end

    blocks.extend(_paragraph_blocks(text[cursor:]))
    return blocks


def _paragraph_alignment(settings: Dict[str, Any]) -> int:
    return _parse_alignment(str(settings.get("paragraphAlign") or "justify")) or WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_body_paragraph(
    doc: Document,
    text: str,
    settings: Optional[Dict[str, Any]] = None,
    indent_first_line: bool = True,
):
    resolved_settings = _normalize_document_settings(settings)
    paragraph = doc.add_paragraph(text if text else " ")
    paragraph.alignment = _paragraph_alignment(resolved_settings)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = float(resolved_settings.get("lineSpacing") or 1.5)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    if paragraph.runs:
        _set_run_font_name(paragraph.runs[0], str(resolved_settings["fontFamily"]))
        paragraph.runs[0].font.size = Pt(float(resolved_settings["bodyFontSize"]))
    return paragraph


def _resolve_image_stream(
    source: str, image_lookup: Optional[Dict[str, Any]] = None
) -> Optional[BytesIO]:
    stripped_source = source.strip()

    if stripped_source.startswith(IMAGE_REF_PREFIX):
        image_id = stripped_source[len(IMAGE_REF_PREFIX) :].strip()
        image_data = image_lookup.get(image_id) if isinstance(image_lookup, dict) else None

        if isinstance(image_data, dict):
            encoded = image_data.get("dataBase64")

            if isinstance(encoded, str) and encoded.strip():
                try:
                    return BytesIO(base64.b64decode(encoded))
                except Exception:
                    return None

    if stripped_source.startswith("data:image/"):
        try:
            _, encoded = stripped_source.split(",", 1)
            return BytesIO(base64.b64decode(encoded))
        except Exception:
            return None

    if stripped_source.startswith(("http://", "https://")):
        try:
            return BytesIO(_fetch_remote_image_bytes(stripped_source))
        except Exception:
            return None

    try:
        local_path = Path(unquote(stripped_source)).expanduser()
        if local_path.is_file():
            return BytesIO(local_path.read_bytes())
    except Exception:
        return None

    return None


@lru_cache(maxsize=32)
def _fetch_remote_image_bytes(url: str) -> bytes:
    response = requests.get(url, timeout=10)
    response.raise_for_status()
    return response.content


def _add_image(
    doc: Document,
    source: str,
    alt: str,
    image_lookup: Optional[Dict[str, Any]] = None,
    max_width_inches: float = 6.0,
    requested_width_inches: Optional[float] = None,
    alignment: str = "center",
    settings: Optional[Dict[str, Any]] = None,
):
    image_stream = _resolve_image_stream(source, image_lookup=image_lookup)
    resolved_settings = _normalize_document_settings(settings)
    paragraph_alignment = _parse_alignment(alignment) or WD_ALIGN_PARAGRAPH.CENTER

    if not image_stream:
        _add_body_paragraph(doc, f"[Image unavailable: {alt}]", resolved_settings)
        return

    try:
        target_width = _clamp_image_width(requested_width_inches, max_width_inches)
        doc.add_picture(image_stream, width=Inches(target_width))
        image_paragraph = doc.paragraphs[-1]
        image_paragraph.alignment = paragraph_alignment
        image_paragraph.paragraph_format.space_after = Pt(6)

        caption = doc.add_paragraph(alt)
        caption.alignment = paragraph_alignment
        caption.paragraph_format.space_after = Pt(10)
        caption.paragraph_format.first_line_indent = Inches(0)

        if caption.runs:
            _set_run_font_name(caption.runs[0], str(resolved_settings["fontFamily"]))
            caption.runs[0].italic = True
            caption.runs[0].font.size = Pt(10)
    except Exception:
        _add_body_paragraph(doc, f"[Image unavailable: {alt}]", resolved_settings)


def _add_logo(
    doc: Document,
    source: str,
    image_lookup: Optional[Dict[str, Any]],
    max_width_inches: float,
    requested_width_inches: Optional[float],
):
    image_stream = _resolve_image_stream(source, image_lookup=image_lookup)
    if not image_stream:
        return

    try:
        target_width = _clamp_image_width(requested_width_inches, max_width_inches)
        doc.add_picture(image_stream, width=Inches(target_width))
        paragraph = doc.paragraphs[-1]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(14)
    except Exception:
        return


def _extract_header_footer_blocks(blocks: List[Dict[str, Any]]):
    header_text = ""
    footer_text = ""
    filtered_blocks: List[Dict[str, Any]] = []

    for block in blocks:
        if not isinstance(block, dict):
            continue

        block_type = str(block.get("type") or "").strip().lower()
        if block_type == "header":
            if not header_text:
                header_text = _extract_plain_text(block.get("html") or block.get("content"), "")
            continue
        if block_type == "footer":
            if not footer_text:
                footer_text = _extract_plain_text(block.get("html") or block.get("content"), "")
            continue

        filtered_blocks.append(block)

    return filtered_blocks, header_text, footer_text


def _set_document_header_footer(doc: Document, header_text: str, footer_text: str, settings: Dict[str, Any]):
    for section in doc.sections:
        if header_text:
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.text = header_text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(10)
                _set_run_font_name(paragraph.runs[0], str(settings["fontFamily"]))

        if footer_text:
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.text = footer_text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(10)
                _set_run_font_name(paragraph.runs[0], str(settings["fontFamily"]))


def _render_title_page(
    doc: Document,
    document_title: str,
    title_page: Optional[Dict[str, Any]],
    image_lookup: Optional[Dict[str, Any]],
    max_width_inches: float,
    settings: Dict[str, Any],
):
    title_page_data = title_page if isinstance(title_page, dict) else {}
    logo_image_id = str(title_page_data.get("logoImageId") or "").strip()
    logo_data_url = str(title_page_data.get("logoDataUrl") or title_page_data.get("logoSrc") or "").strip()
    college_name = _normalized_text(title_page_data.get("collegeName"), "College Name")
    student_name = _normalized_text(title_page_data.get("studentName"), "Student Name")
    course = _normalized_text(title_page_data.get("course") or title_page_data.get("courseName"), "")
    eyebrow = _normalized_text(title_page_data.get("eyebrow"), "")
    subtitle = _normalized_text(title_page_data.get("subtitle"), "")
    note = _normalized_text(title_page_data.get("note"), "")
    logo_width_percent = title_page_data.get("logoWidth")

    requested_logo_width = None
    if isinstance(logo_width_percent, (int, float)):
        requested_logo_width = max_width_inches * max(0.15, min(0.7, float(logo_width_percent) / 100.0))

    logo_source = ""
    if logo_data_url:
        logo_source = logo_data_url
    elif logo_image_id:
        logo_source = f"{IMAGE_REF_PREFIX}{logo_image_id}"

    if logo_source:
        _add_logo(
            doc,
            logo_source,
            image_lookup=image_lookup,
            max_width_inches=max_width_inches,
            requested_width_inches=requested_logo_width,
        )

    if eyebrow:
        eyebrow_paragraph = doc.add_paragraph(eyebrow)
        eyebrow_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eyebrow_paragraph.paragraph_format.space_after = Pt(10)
        if eyebrow_paragraph.runs:
            _set_run_font_name(eyebrow_paragraph.runs[0], str(settings["fontFamily"]))
            eyebrow_paragraph.runs[0].font.size = Pt(max(10, float(settings["bodyFontSize"]) - 1))
            eyebrow_paragraph.runs[0].bold = True

    college = doc.add_paragraph(college_name)
    college.alignment = WD_ALIGN_PARAGRAPH.CENTER
    college.paragraph_format.space_after = Pt(8)
    if college.runs:
        _set_run_font_name(college.runs[0], str(settings["fontFamily"]))
        college.runs[0].font.size = Pt(float(settings["heading3Size"]))
        college.runs[0].bold = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    title = doc.add_paragraph(_normalized_text(document_title, "REPORT"))
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)

    student = doc.add_paragraph(student_name)
    student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if student.runs:
        _set_run_font_name(student.runs[0], str(settings["fontFamily"]))
        student.runs[0].font.size = Pt(float(settings["bodyFontSize"]))

    if course:
        doc.add_paragraph("")
        course_para = doc.add_paragraph(course)
        course_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if course_para.runs:
            _set_run_font_name(course_para.runs[0], str(settings["fontFamily"]))
            course_para.runs[0].font.size = Pt(max(10, float(settings["bodyFontSize"]) - 1))

    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.paragraph_format.space_before = Pt(18)
        subtitle_para.paragraph_format.space_after = Pt(8)
        if subtitle_para.runs:
            _set_run_font_name(subtitle_para.runs[0], str(settings["fontFamily"]))
            subtitle_para.runs[0].font.size = Pt(max(10, float(settings["bodyFontSize"]) - 1))

    if note:
        note_para = doc.add_paragraph(note)
        note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_para.paragraph_format.space_after = Pt(6)
        if note_para.runs:
            _set_run_font_name(note_para.runs[0], str(settings["fontFamily"]))
            note_para.runs[0].font.size = Pt(max(10, float(settings["bodyFontSize"]) - 2))


def _extract_plain_text(value: Any, fallback: str) -> str:
    source = str(value or "").strip()
    if not source:
        return fallback

    if _looks_like_html(source):
        try:
            root = lxml_html.fragment_fromstring(source, create_parent="div")
            text = root.text_content().strip()
            if text:
                return text
        except Exception:
            pass

    normalized = re.sub(r"\s+", " ", source).strip()
    return normalized if normalized else fallback


def _tiptap_escape_html(value: str) -> str:
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _tiptap_marked_text(value: str, marks: List[Dict[str, Any]]) -> str:
    rendered = _tiptap_escape_html(value)
    for mark in marks:
        mark_type = str(mark.get("type") or "").strip().lower()
        if mark_type == "bold":
            rendered = f"<strong>{rendered}</strong>"
        elif mark_type == "italic":
            rendered = f"<em>{rendered}</em>"
        elif mark_type == "underline":
            rendered = f"<u>{rendered}</u>"
        elif mark_type == "strike":
            rendered = f"<s>{rendered}</s>"
        elif mark_type == "code":
            rendered = f"<code>{rendered}</code>"
        elif mark_type == "link":
            attrs = mark.get("attrs") if isinstance(mark.get("attrs"), dict) else {}
            href = _tiptap_escape_html(str(attrs.get("href") or "").strip())
            rendered = f'<a href="{href}">{rendered}</a>' if href else rendered
    return rendered


def _tiptap_inline_html(node: Dict[str, Any]) -> str:
    node_type = str(node.get("type") or "").strip().lower()

    if node_type == "text":
        marks = node.get("marks")
        safe_marks = [mark for mark in marks if isinstance(mark, dict)] if isinstance(marks, list) else []
        return _tiptap_marked_text(str(node.get("text") or ""), safe_marks)

    if node_type == "hardbreak":
        return "<br/>"

    raw_children = node.get("content")
    children = [child for child in raw_children if isinstance(child, dict)] if isinstance(raw_children, list) else []
    return "".join(_tiptap_inline_html(child) for child in children)


def _tiptap_plain_text(node: Dict[str, Any]) -> str:
    node_type = str(node.get("type") or "").strip().lower()

    if node_type == "text":
        return str(node.get("text") or "")

    if node_type == "hardbreak":
        return "\n"

    raw_children = node.get("content")
    children = [child for child in raw_children if isinstance(child, dict)] if isinstance(raw_children, list) else []
    return "".join(_tiptap_plain_text(child) for child in children)


def _tiptap_alignment_style(node: Dict[str, Any]) -> str:
    attrs = node.get("attrs") if isinstance(node.get("attrs"), dict) else {}
    align = str(attrs.get("textAlign") or attrs.get("align") or "").strip().lower()
    if align in {"left", "center", "right", "justify"}:
        return f' style="text-align:{align}"'
    return ""


def _tiptap_list_html(node: Dict[str, Any], ordered: bool) -> str:
    tag = "ol" if ordered else "ul"
    raw_items = node.get("content")
    items = [item for item in raw_items if isinstance(item, dict)] if isinstance(raw_items, list) else []
    if not items:
        return f"<{tag}></{tag}>"

    rendered_items: List[str] = []
    for item in items:
        if str(item.get("type") or "").strip().lower() != "listitem":
            continue
        raw_children = item.get("content")
        children = [child for child in raw_children if isinstance(child, dict)] if isinstance(raw_children, list) else []

        inline_parts: List[str] = []
        nested_parts: List[str] = []

        for child in children:
            child_type = str(child.get("type") or "").strip().lower()
            if child_type == "orderedlist":
                nested_parts.append(_tiptap_list_html(child, ordered=True))
            elif child_type == "bulletlist":
                nested_parts.append(_tiptap_list_html(child, ordered=False))
            else:
                inline_parts.append(_tiptap_inline_html(child))

        item_content = "".join(inline_parts).strip() or "&nbsp;"
        rendered_items.append(f"<li>{item_content}{''.join(nested_parts)}</li>")

    return f"<{tag}>{''.join(rendered_items)}</{tag}>"


def _tiptap_table_rows(node: Dict[str, Any]) -> List[List[str]]:
    raw_rows = node.get("content")
    rows = [row for row in raw_rows if isinstance(row, dict)] if isinstance(raw_rows, list) else []
    result: List[List[str]] = []

    for row in rows:
        if str(row.get("type") or "").strip().lower() != "tablerow":
            continue
        raw_cells = row.get("content")
        cells = [cell for cell in raw_cells if isinstance(cell, dict)] if isinstance(raw_cells, list) else []

        row_values: List[str] = []
        for cell in cells:
            cell_type = str(cell.get("type") or "").strip().lower()
            if cell_type not in {"tablecell", "tableheader"}:
                continue
            row_values.append(_tiptap_plain_text(cell).strip())

        result.append(row_values)

    return result


def tiptap_document_to_blocks(document_data: Any) -> List[Dict[str, Any]]:
    if not isinstance(document_data, dict):
        return []

    root_type = str(document_data.get("type") or "").strip().lower()
    if root_type and root_type != "doc":
        return []

    raw_nodes = document_data.get("content")
    nodes = [node for node in raw_nodes if isinstance(node, dict)] if isinstance(raw_nodes, list) else []

    blocks: List[Dict[str, Any]] = []
    for node in nodes:
        node_type = str(node.get("type") or "").strip().lower()
        alignment = _tiptap_alignment_style(node)

        if node_type == "heading":
            attrs = node.get("attrs") if isinstance(node.get("attrs"), dict) else {}
            level = int(attrs.get("level") or 1)
            level = min(max(level, 1), 3)
            html_value = _tiptap_inline_html(node).strip() or "Untitled Heading"
            blocks.append({"type": f"heading{level}", "html": html_value})
            continue

        if node_type == "paragraph":
            html_value = _tiptap_inline_html(node)
            blocks.append({"type": "paragraph", "html": f"<p{alignment}>{html_value}</p>" if html_value else "<p></p>"})
            continue

        if node_type == "bulletlist":
            blocks.append({"type": "bullet_list", "html": _tiptap_list_html(node, ordered=False)})
            continue

        if node_type == "orderedlist":
            blocks.append({"type": "numbered_list", "html": _tiptap_list_html(node, ordered=True)})
            continue

        if node_type == "blockquote":
            html_value = _tiptap_inline_html(node)
            blocks.append({"type": "quote", "html": f"<blockquote>{html_value}</blockquote>" if html_value else "<blockquote></blockquote>"})
            continue

        if node_type == "codeblock":
            blocks.append({"type": "code", "code": _tiptap_plain_text(node)})
            continue

        if node_type == "table":
            rows = _tiptap_table_rows(node)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        if node_type == "image":
            attrs = node.get("attrs") if isinstance(node.get("attrs"), dict) else {}
            src = str(attrs.get("src") or "").strip()
            if not src:
                continue
            caption = str(attrs.get("alt") or attrs.get("title") or "Figure").strip() or "Figure"
            width_attr = attrs.get("width")
            width = 75
            if isinstance(width_attr, str):
                match = PERCENT_PATTERN.search(width_attr)
                if match:
                    width = int(float(match.group(1)))
            elif isinstance(width_attr, (int, float)):
                width = int(float(width_attr))
            width = max(20, min(100, width))
            blocks.append({"type": "image", "source": src, "caption": caption, "width": width})
            continue

        if node_type in {"horizontalrule", "pagebreak"}:
            blocks.append({"type": "page_break"})
            continue

        fallback = _tiptap_plain_text(node).strip()
        blocks.append({"type": "paragraph", "html": f"<p>{_tiptap_escape_html(fallback)}</p>" if fallback else "<p></p>"})

    return blocks


def _render_code_block(doc: Document, code_text: str, settings: Optional[Dict[str, Any]] = None):
    lines = str(code_text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if not lines:
        lines = [""]

    for line in lines:
        paragraph = doc.add_paragraph(line if line else " ")
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.right_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(" ")
        _set_run_font_name(run, "Courier New")
        run.font.size = Pt(10)


def _render_table_block(doc: Document, rows: Any, settings: Optional[Dict[str, Any]] = None):
    resolved_settings = _normalize_document_settings(settings)
    if not isinstance(rows, list) or not rows:
        _add_body_paragraph(doc, "", resolved_settings)
        return

    normalized_rows: List[List[str]] = []
    max_cols = 0

    for raw_row in rows:
        if not isinstance(raw_row, list):
            continue

        row_values = [str(cell or "") for cell in raw_row]
        max_cols = max(max_cols, len(row_values))
        normalized_rows.append(row_values)

    if not normalized_rows or max_cols <= 0:
        _add_body_paragraph(doc, "", resolved_settings)
        return

    table = doc.add_table(rows=len(normalized_rows), cols=max_cols)
    table.style = "Table Grid"

    for row_index, row_values in enumerate(normalized_rows):
        for col_index in range(max_cols):
            cell_text = row_values[col_index] if col_index < len(row_values) else ""
            cell = table.rows[row_index].cells[col_index]
            cell.text = cell_text

            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _set_run_font_name(run, str(resolved_settings["fontFamily"]))
                    run.font.size = Pt(max(10, float(resolved_settings["bodyFontSize"]) - 1))

    doc.add_paragraph("")


def _sections_to_blocks(sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    blocks: List[Dict[str, Any]] = []

    for section in sections:
        if not isinstance(section, dict):
            continue

        section_title = _normalized_text(section.get("title"), "Untitled Section")
        section_content = _normalized_text(section.get("content"), "")
        blocks.append({"type": "heading1", "html": section_title})
        blocks.append({"type": "paragraph", "html": section_content})

        raw_subsections = section.get("subsections", [])
        subsections = raw_subsections if isinstance(raw_subsections, list) else []

        for subsection in subsections:
            if not isinstance(subsection, dict):
                continue

            subsection_title = _normalized_text(subsection.get("title"), "Untitled Subsection")
            subsection_content = _normalized_text(subsection.get("content"), "")
            blocks.append({"type": "heading2", "html": subsection_title})
            blocks.append({"type": "paragraph", "html": subsection_content})

    return blocks


def _replace_reference_and_footnote_tokens(
    value: str,
    reference_lookup: Dict[str, Dict[str, Any]],
    footnote_lookup: Dict[str, Dict[str, Any]],
) -> str:
    source = str(value or "")

    def replace_citation(match):
        key = match.group(1)
        item = reference_lookup.get(key)
        return f"[{item['index']}]" if item else match.group(0)

    def replace_footnote(match):
        key = match.group(1)
        item = footnote_lookup.get(key)
        return f"({item['index']})" if item else match.group(0)

    return FOOTNOTE_PATTERN.sub(replace_footnote, CITATION_PATTERN.sub(replace_citation, source))


def _collect_reference_lookup(blocks: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    lookup: Dict[str, Dict[str, Any]] = {}
    index = 1

    for block in blocks:
        if not isinstance(block, dict):
            continue
        if str(block.get("type") or "").strip().lower() != "reference":
            continue

        key = str(block.get("citationKey") or block.get("citation_key") or "").strip()
        if not key or key in lookup:
            continue

        source = str(block.get("source") or "").strip()
        lookup[key] = {
            "index": index,
            "source": source if source else key,
        }
        index += 1

    return lookup


def _collect_footnote_lookup(blocks: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    lookup: Dict[str, Dict[str, Any]] = {}
    index = 1

    for block in blocks:
        if not isinstance(block, dict):
            continue
        if str(block.get("type") or "").strip().lower() != "footnote":
            continue

        key = str(block.get("footnoteKey") or block.get("footnote_key") or "").strip()
        if not key or key in lookup:
            continue

        content = str(block.get("content") or "").strip()
        lookup[key] = {
            "index": index,
            "content": content if content else key,
        }
        index += 1

    return lookup


def _render_equation_block(doc: Document, latex: str, label: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)

    equation_run = paragraph.add_run(latex if latex else "E = mc^2")
    _set_run_font_name(equation_run, "Cambria Math")
    equation_run.font.size = Pt(12)

    if label:
        paragraph.add_run(" ")
        label_run = paragraph.add_run(f"({label})")
        _set_run_font_name(label_run, "Times New Roman")
        label_run.font.size = Pt(10)
        label_run.italic = True


def _append_reference_section(doc: Document, reference_lookup: Dict[str, Dict[str, Any]]):
    if not reference_lookup:
        return

    doc.add_paragraph("")
    heading = doc.add_heading("References", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    ordered_items = sorted(reference_lookup.items(), key=lambda item: item[1]["index"])
    for _, reference in ordered_items:
        paragraph = doc.add_paragraph(f"[{reference['index']}] {reference['source']}")
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _append_footnote_section(doc: Document, footnote_lookup: Dict[str, Dict[str, Any]]):
    if not footnote_lookup:
        return

    doc.add_paragraph("")
    heading = doc.add_heading("Footnotes", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    ordered_items = sorted(footnote_lookup.items(), key=lambda item: item[1]["index"])
    for _, footnote in ordered_items:
        paragraph = doc.add_paragraph(f"{footnote['index']}. {footnote['content']}")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if paragraph.runs:
            paragraph.runs[0].font.size = Pt(10)


def _append_comment_section(doc: Document, comments: Optional[List[Dict[str, Any]]]):
    if not comments:
        return

    normalized_comments = [comment for comment in comments if isinstance(comment, dict)]
    if not normalized_comments:
        return

    doc.add_paragraph("")
    heading = doc.add_heading("Comments", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for index, comment in enumerate(normalized_comments, start=1):
        text = str(comment.get("text") or "").strip()
        if not text:
            continue

        block_id = str(comment.get("blockId") or comment.get("block_id") or "").strip()
        author = str(comment.get("author") or "Editor").strip()
        status = "Resolved" if bool(comment.get("resolved")) else "Open"
        prefix = f"{index}. [{status}] {author}"
        if block_id:
            prefix = f"{prefix} (Block {block_id})"

        paragraph = doc.add_paragraph(f"{prefix}: {text}")
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _render_document_blocks(
    doc: Document,
    blocks: List[Dict[str, Any]],
    image_lookup: Optional[Dict[str, Any]],
    max_width_inches: float,
    comments: Optional[List[Dict[str, Any]]] = None,
    settings: Optional[Dict[str, Any]] = None,
):
    resolved_settings = _normalize_document_settings(settings)
    heading_1 = 0
    heading_2 = 0
    heading_3 = 0
    rendered_content_blocks = 0
    reference_lookup = _collect_reference_lookup(blocks)
    footnote_lookup = _collect_footnote_lookup(blocks)

    for raw_block in blocks:
        if not isinstance(raw_block, dict):
            continue

        block_type = str(raw_block.get("type") or "paragraph").strip().lower()

        if block_type == "page_break":
            doc.add_page_break()
            continue

        if block_type == "heading1":
            if rendered_content_blocks > 0 and resolved_settings.get("pageBreakAfterHeading1"):
                doc.add_page_break()
            title_text = _extract_plain_text(
                raw_block.get("html") or raw_block.get("content") or raw_block.get("title"),
                "Untitled Section",
            )
            title_text = _replace_reference_and_footnote_tokens(
                title_text,
                reference_lookup,
                footnote_lookup,
            )
            heading_1 += 1
            heading_2 = 0
            heading_3 = 0

            heading = doc.add_heading(f"{heading_1}. {title_text}", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            rendered_content_blocks += 1
            continue

        if block_type == "heading2":
            title_text = _extract_plain_text(
                raw_block.get("html") or raw_block.get("content") or raw_block.get("title"),
                "Untitled Subsection",
            )
            title_text = _replace_reference_and_footnote_tokens(
                title_text,
                reference_lookup,
                footnote_lookup,
            )
            if heading_1 == 0:
                heading_1 = 1
            heading_2 += 1
            heading_3 = 0

            heading = doc.add_heading(f"{heading_1}.{heading_2}. {title_text}", level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            rendered_content_blocks += 1
            continue

        if block_type == "heading3":
            title_text = _extract_plain_text(
                raw_block.get("html") or raw_block.get("content") or raw_block.get("title"),
                "Untitled Subheading",
            )
            title_text = _replace_reference_and_footnote_tokens(
                title_text,
                reference_lookup,
                footnote_lookup,
            )
            if heading_1 == 0:
                heading_1 = 1
            if heading_2 == 0:
                heading_2 = 1
            heading_3 += 1

            heading = doc.add_heading(f"{heading_1}.{heading_2}.{heading_3}. {title_text}", level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            rendered_content_blocks += 1
            continue

        if block_type in {"paragraph", "bullet_list", "numbered_list"}:
            content = _replace_reference_and_footnote_tokens(
                str(raw_block.get("html") or raw_block.get("content") or ""),
                reference_lookup,
                footnote_lookup,
            )
            _render_content(
                doc,
                content,
                image_lookup=image_lookup,
                max_width_inches=max_width_inches,
                settings=resolved_settings,
            )
            rendered_content_blocks += 1
            continue

        if block_type == "quote":
            quote_text = _extract_plain_text(raw_block.get("html") or raw_block.get("content"), " ")
            quote_text = _replace_reference_and_footnote_tokens(
                quote_text,
                reference_lookup,
                footnote_lookup,
            )
            quote_paragraph = doc.add_paragraph(quote_text)
            quote_paragraph.alignment = _paragraph_alignment(resolved_settings)
            quote_paragraph.paragraph_format.left_indent = Inches(0.35)
            quote_paragraph.paragraph_format.right_indent = Inches(0.2)
            quote_paragraph.paragraph_format.space_after = Pt(8)
            quote_paragraph.paragraph_format.line_spacing = float(resolved_settings["lineSpacing"])

            if quote_paragraph.runs:
                _set_run_font_name(quote_paragraph.runs[0], str(resolved_settings["fontFamily"]))
                quote_paragraph.runs[0].font.size = Pt(float(resolved_settings["bodyFontSize"]))
                quote_paragraph.runs[0].italic = True
            rendered_content_blocks += 1
            continue

        if block_type == "code":
            _render_code_block(doc, str(raw_block.get("code") or ""), resolved_settings)
            doc.add_paragraph("")
            rendered_content_blocks += 1
            continue

        if block_type == "table":
            _render_table_block(doc, raw_block.get("rows"), resolved_settings)
            rendered_content_blocks += 1
            continue

        if block_type == "equation":
            latex = str(raw_block.get("latex") or "").strip()
            label = str(raw_block.get("label") or "").strip()
            _render_equation_block(doc, latex, label)
            rendered_content_blocks += 1
            continue

        if block_type in {"reference", "footnote", "header", "footer"}:
            continue

        if block_type == "image":
            image_id = str(raw_block.get("imageId") or raw_block.get("image_id") or "").strip()
            image_source = str(raw_block.get("source") or "").strip()
            image_alignment = str(raw_block.get("alignment") or "center").strip().lower()

            if image_id:
                image_source = f"{IMAGE_REF_PREFIX}{image_id}"

            if not image_source:
                _add_body_paragraph(doc, "[Image unavailable: Figure]")
                continue

            caption = _normalized_text(raw_block.get("caption"), "Figure")
            width_value = raw_block.get("width")
            requested_width: Optional[float] = None

            if isinstance(width_value, (int, float)):
                width_ratio = max(0.1, min(1.0, float(width_value) / 100.0))
                requested_width = max_width_inches * width_ratio

            _add_image(
                doc,
                image_source,
                caption,
                image_lookup=image_lookup,
                max_width_inches=max_width_inches,
                requested_width_inches=requested_width,
                alignment=image_alignment,
                settings=resolved_settings,
            )
            rendered_content_blocks += 1
            continue

        fallback_content = str(raw_block.get("html") or raw_block.get("content") or raw_block.get("text") or "")
        fallback_content = _replace_reference_and_footnote_tokens(
            fallback_content,
            reference_lookup,
            footnote_lookup,
        )
        _render_content(
            doc,
            fallback_content,
            image_lookup=image_lookup,
            max_width_inches=max_width_inches,
            settings=resolved_settings,
        )
        rendered_content_blocks += 1

    _append_reference_section(doc, reference_lookup)
    _append_footnote_section(doc, footnote_lookup)
    _append_comment_section(doc, comments)


def _merge_inline_style(inherited_style: Dict[str, Any], element) -> Dict[str, Any]:
    next_style = dict(inherited_style)
    tag = str(element.tag).lower()
    style_map = _parse_style_map(element.get("style", ""))

    if tag in {"strong", "b"}:
        next_style["bold"] = True
    if tag in {"em", "i"}:
        next_style["italic"] = True
    if tag == "u":
        next_style["underline"] = True
    if tag == "mark":
        next_style["highlight"] = WD_COLOR_INDEX.YELLOW

    if style_map.get("font-weight", "").strip() in {"700", "800", "900", "bold"}:
        next_style["bold"] = True
    if style_map.get("font-style", "").strip() == "italic":
        next_style["italic"] = True
    if "underline" in style_map.get("text-decoration", "").strip():
        next_style["underline"] = True

    font_name = style_map.get("font-family")
    if font_name:
        next_style["font_name"] = font_name.split(",")[0].strip().strip('"').strip("'")

    font_size_pt = _parse_font_size_pt(style_map.get("font-size"))
    if font_size_pt is not None:
        next_style["font_size_pt"] = font_size_pt

    highlight = _parse_highlight_color(style_map.get("background-color", ""))
    if highlight is not None:
        next_style["highlight"] = highlight

    return next_style


def _append_text_with_style(paragraph, text: str, style: Dict[str, Any]):
    if not text:
        return

    run = paragraph.add_run(text)
    _apply_inline_run_style(run, style)


def _render_inline_nodes(paragraph, element, inherited_style: Dict[str, Any]):
    if element.text:
        _append_text_with_style(paragraph, element.text, inherited_style)

    for child in element:
        tag = str(child.tag).lower()

        if tag == "br":
            _append_text_with_style(paragraph, "\n", inherited_style)
        elif tag == "img":
            continue
        else:
            child_style = _merge_inline_style(inherited_style, child)
            _render_inline_nodes(paragraph, child, child_style)

        if child.tail:
            _append_text_with_style(paragraph, child.tail, inherited_style)


def _render_list_block(doc: Document, block, ordered: bool, settings: Optional[Dict[str, Any]] = None):
    resolved_settings = _normalize_document_settings(settings)
    parent_style_map = _parse_style_map(block.get("style", ""))
    parent_alignment = _extract_alignment(block, parent_style_map)

    for item in block.xpath("./li"):
        paragraph = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = float(resolved_settings["lineSpacing"])

        item_styles = _parse_style_map(item.get("style", ""))
        item_alignment = _extract_alignment(item, item_styles)
        paragraph.alignment = item_alignment if item_alignment is not None else parent_alignment

        base_style = {
            "font_name": str(resolved_settings["fontFamily"]),
            "font_size_pt": float(resolved_settings["bodyFontSize"]),
        }
        _render_inline_nodes(paragraph, item, base_style)

        if not paragraph.runs:
            paragraph.add_run(" ")


def _render_html_block(
    doc: Document,
    block,
    image_lookup: Optional[Dict[str, Any]],
    max_width_inches: float,
    settings: Optional[Dict[str, Any]] = None,
):
    resolved_settings = _normalize_document_settings(settings)
    tag = str(block.tag).lower()

    if tag in {"ul", "ol"}:
        _render_list_block(doc, block, ordered=(tag == "ol"), settings=resolved_settings)
        return

    if tag in {"figure", "img"}:
        image_element = block if tag == "img" else block.xpath(".//img[1]")
        if isinstance(image_element, list):
            image_element = image_element[0] if image_element else None

        if image_element is None:
            return

        style_map = _parse_style_map(image_element.get("style", ""))
        width_value = style_map.get("width") or image_element.get("width", "")
        requested_width = _parse_width_inches(width_value, max_width_inches)
        alignment_value = style_map.get("text-align") or block.get("align") or "center"

        caption_text = (image_element.get("alt") or "Figure").strip() or "Figure"
        if tag == "figure":
            caption_nodes = block.xpath(".//figcaption")
            if caption_nodes:
                caption = caption_nodes[0].text_content().strip()
                if caption:
                    caption_text = caption

        _add_image(
            doc,
            image_element.get("src", ""),
            caption_text,
            image_lookup=image_lookup,
            max_width_inches=max_width_inches,
            requested_width_inches=requested_width,
            alignment=str(alignment_value),
            settings=resolved_settings,
        )
        return

    if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        heading_level = int(tag[1]) if tag[1].isdigit() else 2
        heading_level = min(max(1, heading_level), 3)
        heading = doc.add_heading("", level=heading_level)
        heading.paragraph_format.space_after = Pt(8)
        _render_inline_nodes(
            heading,
            block,
            {
                "font_name": str(resolved_settings["fontFamily"]),
                "font_size_pt": float(resolved_settings[f"heading{heading_level}Size"]),
            },
        )
        if not heading.runs:
            heading.add_run(" ")
        return

    paragraph = _add_body_paragraph(doc, "", resolved_settings, indent_first_line=True)
    style_map = _parse_style_map(block.get("style", ""))
    alignment = _extract_alignment(block, style_map)
    if alignment is not None:
        paragraph.alignment = alignment

    base_style = {
        "font_name": str(resolved_settings["fontFamily"]),
        "font_size_pt": float(resolved_settings["bodyFontSize"]),
    }
    _render_inline_nodes(paragraph, block, base_style)
    if not paragraph.runs:
        paragraph.add_run(" ")


def _render_html_content(
    doc: Document,
    content: str,
    image_lookup: Optional[Dict[str, Any]],
    max_width_inches: float,
    settings: Optional[Dict[str, Any]] = None,
) -> bool:
    source = str(content or "").strip()

    if not source:
        return False

    try:
        root = lxml_html.fragment_fromstring(source, create_parent="div")
    except Exception:
        return False

    consumed = False

    if root.text and root.text.strip():
        _add_body_paragraph(doc, root.text.strip(), settings)
        consumed = True

    for block in root:
        _render_html_block(
            doc,
            block,
            image_lookup=image_lookup,
            max_width_inches=max_width_inches,
            settings=settings,
        )
        consumed = True

        if block.tail and block.tail.strip():
            _add_body_paragraph(doc, block.tail.strip(), settings)

    return consumed


def _render_legacy_content(
    doc: Document,
    content: str,
    image_lookup: Optional[Dict[str, Any]] = None,
    max_width_inches: float = 6.0,
    settings: Optional[Dict[str, Any]] = None,
):
    blocks = _content_blocks(content)

    if not blocks:
        _add_body_paragraph(doc, "", settings)
        return

    for block in blocks:
        if block["type"] == "paragraph":
            _add_body_paragraph(doc, block["text"], settings)
        elif block["type"] == "image":
            _add_image(
                doc,
                block["source"],
                block["alt"],
                image_lookup=image_lookup,
                max_width_inches=max_width_inches,
                settings=settings,
            )


def _render_content(
    doc: Document,
    content: str,
    image_lookup: Optional[Dict[str, Any]] = None,
    max_width_inches: float = 6.0,
    settings: Optional[Dict[str, Any]] = None,
):
    normalized_content = str(content or "")

    if _looks_like_html(normalized_content):
        rendered = _render_html_content(
            doc,
            normalized_content,
            image_lookup=image_lookup,
            max_width_inches=max_width_inches,
            settings=settings,
        )
        if rendered:
            return

    _render_legacy_content(
        doc,
        normalized_content,
        image_lookup=image_lookup,
        max_width_inches=max_width_inches,
        settings=settings,
    )


def add_table_of_contents(doc: Document, heading_levels: str = "1-2"):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")

    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = f'TOC \\o "{heading_levels}" \\h \\z \\u'

    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")

    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "Right-click and update field to build the Table of Contents."

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    run._r.append(field_begin)
    run._r.append(instruction_text)
    run._r.append(field_separate)
    run._r.append(placeholder_text)
    run._r.append(field_end)


def _add_page_number_field(paragraph):
    run = paragraph.add_run()

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")

    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = "PAGE"

    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    run._r.append(field_begin)
    run._r.append(instruction_text)
    run._r.append(field_separate)
    run._r.append(field_end)


def _apply_page_numbers(doc: Document):
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.text.strip():
            paragraph.add_run("   ")
        _add_page_number_field(paragraph)


def create_editor_docx(
    output_path,
    document_title: str = "REPORT",
    image_lookup: Optional[Dict[str, Any]] = None,
    sections: Optional[List[Dict[str, Any]]] = None,
    blocks: Optional[List[Dict[str, Any]]] = None,
    comments: Optional[List[Dict[str, Any]]] = None,
    title_page: Optional[Dict[str, Any]] = None,
    document_settings: Optional[Dict[str, Any]] = None,
    document_structure: Optional[Dict[str, Any]] = None,
):
    settings = _normalize_document_settings(document_settings)
    structure = _normalize_document_structure(document_structure)
    doc = Document()
    _configure_document(doc, settings)
    page = doc.sections[0]
    max_image_width = (
        page.page_width.inches - page.left_margin.inches - page.right_margin.inches
    )

    valid_blocks = [block for block in (blocks or []) if isinstance(block, dict)]
    if not valid_blocks:
        valid_sections = [section for section in (sections or []) if isinstance(section, dict)]
        valid_blocks = _sections_to_blocks(valid_sections)

    if not valid_blocks:
        valid_blocks = [{"type": "paragraph", "html": ""}]

    valid_blocks, header_text, footer_text = _extract_header_footer_blocks(valid_blocks)

    if structure["showCoverPage"]:
        _render_title_page(
            doc,
            document_title=_normalized_text(document_title, "REPORT"),
            title_page=title_page,
            image_lookup=image_lookup,
            max_width_inches=max_image_width,
            settings=settings,
        )

    if structure["showTableOfContents"]:
        if doc.paragraphs:
            doc.add_page_break()
        toc_heading = doc.add_heading("Table of Contents", level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table_of_contents(doc, heading_levels="1-3")

    if isinstance(title_page, dict):
        if not header_text:
            header_text = _extract_plain_text(title_page.get("headerText"), "")
        if not footer_text:
            footer_text = _extract_plain_text(title_page.get("footerText"), "")

    _set_document_header_footer(doc, header_text, footer_text, settings)
    _apply_page_numbers(doc)

    if structure["showCoverPage"] or structure["showTableOfContents"]:
        doc.add_page_break()

    _render_document_blocks(
        doc,
        valid_blocks,
        image_lookup=image_lookup,
        max_width_inches=max_image_width,
        comments=comments,
        settings=settings,
    )

    doc.save(output_path)


def create_professional_docx(structured_data, output_path):
    doc = Document()
    _configure_document(doc, _normalize_document_settings())
    _apply_page_numbers(doc)

    title = doc.add_paragraph("LAB REPORT")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    toc_title = doc.add_paragraph("Table of Contents")
    toc_title.style = "Heading 1"
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table_of_contents(doc, heading_levels="1-3")
    doc.add_page_break()

    for section_data in structured_data:
        heading = doc.add_heading(section_data["heading"], level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _render_content(doc, section_data.get("content", ""))

    doc.save(output_path)
